from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QLabel, QPushButton, QTableWidget, QTableWidgetItem,
    QHBoxLayout, QLineEdit, QTextEdit, QDateEdit, QMessageBox, QFormLayout, QFileDialog,
    QHeaderView, QDialog
)
from PyQt5.QtCore import Qt, QDate
import psycopg2
from docx import Document
from datetime import datetime

def get_db_connection():
    return psycopg2.connect(
        host="localhost",
        database="MediTrackSNHS",
        user="postgres",
        password="Mylovemondejar"
    )

class ReportsPage(QWidget):
    def __init__(self, conn=None):
        super().__init__()
        self.conn = conn
        self.records_page = ReportsRecordsPage(self.conn)
        layout = QVBoxLayout(self)
        layout.addWidget(self.records_page)
        self.setLayout(layout)
        self.setStyleSheet("""
            QWidget {
                background:#FFE1D3
            }
        """)

class ReportsRecordsPage(QWidget):
    def __init__(self, conn=None):
        super().__init__()
        self.conn = conn
        self.init_ui()

    def get_conn(self):
        if self.conn:
            return self.conn
        else:
            return get_db_connection()

    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 60)
        layout.setSpacing(10)

        title = QLabel("Accident/Incident Reports")
        title.setStyleSheet("""
            font-size: 28px;
            font-weight: bold;
            background: #FDD1B0;
            border-radius: 10px;
            padding: 10px;
            color: #3a2b23;
        """)
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)

        # Controls
        control_layout = QHBoxLayout()
        self.search_bar = QLineEdit()
        self.search_bar.setFixedHeight(45)
        self.search_bar.setFixedWidth(400)
        self.search_bar.setStyleSheet("""
            QLineEdit {
                background-color: #FFF0F5;
                border: 2px solid #D295BE;
                border-radius: 5px;
                padding: 8px 12px;
                font-size: 16px;
                color: #333;
            }
            QLineEdit:focus {
                border: 2px solid #1a5276;
                background-color: #eaf2f8;
            }
        """)
        self.search_bar.setPlaceholderText("Search report...")
        self.search_bar.textChanged.connect(self.search_reports)

        add_btn = QPushButton("Add Report")
        add_btn.setFixedWidth(150)
        add_btn.setStyleSheet("background-color: #D295BE; color: white; font-weight: bold; padding: 15px; border-radius: 10px; font-size: 18px;")
        add_btn.clicked.connect(self.open_add_report_dialog)

        self.edit_btn = QPushButton("Edit")
        self.edit_btn.setFixedWidth(120)
        self.edit_btn.setStyleSheet("background-color: #f7ca18; color: white; font-weight: bold; padding: 15px; border-radius: 10px; font-size: 18px;")
        self.edit_btn.clicked.connect(self.edit_record)

        self.view_btn = QPushButton("View")
        self.view_btn.setFixedWidth(120)
        self.view_btn.setStyleSheet("background-color: #5dade2; color: white; font-weight: bold; padding: 15px; border-radius: 10px; font-size: 18px;")
        self.view_btn.clicked.connect(self.view_record)

        self.delete_btn = QPushButton("Delete")
        self.delete_btn.setFixedWidth(120)
        self.delete_btn.setStyleSheet("background-color: #ff9999; color: white; font-weight: bold; padding: 15px; border-radius: 10px; font-size: 18px;")
        self.delete_btn.clicked.connect(self.delete_record)

        control_layout.addWidget(self.search_bar)
        control_layout.addStretch()
        control_layout.addWidget(add_btn)
        control_layout.addWidget(self.edit_btn)
        control_layout.addWidget(self.view_btn)
        control_layout.addWidget(self.delete_btn)
        layout.addLayout(control_layout)

        # Table -- include Student Name columns
        self.table = QTableWidget()
        self.table.setColumnCount(11)
        self.table.setHorizontalHeaderLabels([
            "Student LRN", "First Name", "Last Name", "Middle Name", "Date", "Description",
            "Action Taken", "Referral", "Follow-up Date", "Created At", "Actions"
        ])
        self.table.setStyleSheet("""
            QTableWidget {
                background-color: #FFF5F0;
            }
            QTableWidget::item {
                font-size: 18px;
            }
            QHeaderView::section {
                font-size: 20px;
                font-weight: bold;
                background-color: #FFE1D3;
                border: 1px solid #D295BE;
            }
            QTableCornerButton::section {
                background-color: #FFE1D3;
                border: 1px solid #D295BE;
            }
        """)
        self.table.setSelectionBehavior(QTableWidget.SelectRows)
        self.table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.table.setSortingEnabled(False)
        header = self.table.horizontalHeader()
        for i in range(self.table.columnCount()):
            header.setSectionResizeMode(i, QHeaderView.Stretch)
        self.table.horizontalHeader().setStretchLastSection(True)
        layout.addWidget(self.table)
        self.setLayout(layout)
        self.load_data()

    def load_data(self):
        self.table.setRowCount(0)
        try:
            conn = self.get_conn()
            cur = conn.cursor()
            cur.execute("""
                SELECT ir.ir_id, ir.stud_id, s.stud_fname, s.stud_lname, s.stud_mname, 
                       ir.ir_date_incident, ir.ir_description, ir.ir_action_taken,
                       ir.ir_referral, ir.ir_follow_up_date, ir.created_at
                FROM incident_report ir
                LEFT JOIN student s ON ir.stud_id = s.stud_id
                ORDER BY ir.ir_date_incident DESC
            """)
            rows = cur.fetchall()
            self.rows = rows
            self.display_rows(self.rows)
            cur.close()
            if not self.conn:
                conn.close()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load reports: {e}")

    def display_rows(self, rows):
        self.table.setRowCount(len(rows))
        for row_idx, row in enumerate(rows):
            # row = (ir_id, stud_id, fname, lname, mname, date, desc, action, referral, followup, created_at)
            for col_idx, val in enumerate(row[1:]):  # skip ir_id
                item = QTableWidgetItem(str(val) if val is not None else "")
                font = item.font()
                font.setPointSize(9)
                item.setFont(font)
                self.table.setItem(row_idx, col_idx, item)
            self.table.setVerticalHeaderItem(row_idx, QTableWidgetItem(str(row_idx + 1)))
            btn = QPushButton("Generate DOCX")
            btn.setFixedWidth(136)
            btn.setStyleSheet("""
                QPushButton { background: #A3D5FF; border-radius:5px; font-size:16px; padding:6px 12px;}
            """)
            btn.clicked.connect(lambda _, r=row: self.generate_report_docx(r))
            btn_layout = QHBoxLayout()
            btn_layout.addWidget(btn)
            btn_layout.setContentsMargins(0,0,0,0)
            btn_widget = QWidget()
            btn_widget.setStyleSheet("background-color: #f6fbfc;")
            btn_widget.setLayout(btn_layout)
            self.table.setCellWidget(row_idx, 10, btn_widget)

    def search_reports(self):
        query = self.search_bar.text().lower()
        filtered = []
        for row in self.rows:
            if any(query in str(col).lower() for col in row[1:]):
                filtered.append(row)
        self.display_rows(filtered)

    def generate_report_docx(self, report_row):
        report = {
            "stud_id": report_row[1],
            "fname": report_row[2],
            "lname": report_row[3],
            "mname": report_row[4],
            "date": str(report_row[5]),
            "description": report_row[6],
            "action_taken": report_row[7],
            "referral": report_row[8],
            "follow_up_date": str(report_row[9]),
            "created_at": str(report_row[10])
        }
        default_filename = f"IncidentReport_{report['stud_id']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        save_path, _ = QFileDialog.getSaveFileName(self, "Save Report DOCX", default_filename, "Word Documents (*.docx)")
        if not save_path:
            return
        try:
            doc = Document()
            doc.add_heading('Accident/Incident Report', 0)
            doc.add_paragraph(f"Student LRN: {report['stud_id']}")
            doc.add_paragraph(f"Name: {report['fname']} {report['mname']} {report['lname']}")
            doc.add_paragraph(f"Date of Incident: {report['date']}")
            doc.add_paragraph(f"Description:\n{report['description']}")
            doc.add_paragraph(f"Action Taken:\n{report['action_taken']}")
            doc.add_paragraph(f"Referral: {report['referral']}")
            doc.add_paragraph(f"Follow-up Date: {report['follow_up_date']}")
            doc.add_paragraph(f"Report Created At: {report['created_at']}")
            doc.add_paragraph("\nSignature: ___________________________")
            doc.add_paragraph("Clinic Staff")
            doc.save(save_path)
            QMessageBox.information(self, "Report Doc Created", f"Report document saved as:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "DOCX Error", f"Could not generate report document: {e}")

    def open_add_report_dialog(self):
        dialog = AddReportDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            data = dialog.get_data()
            if data:
                lrn, date, desc, action, referral, followup = data

                def empty_to_none(x):
                    return x if x.strip() != "" else None

                lrn = lrn.strip()
                date = date.strip()
                desc = desc.strip()
                action = empty_to_none(action)
                referral = empty_to_none(referral)
                followup = empty_to_none(followup)

                try:
                    conn = self.get_conn()
                    cur = conn.cursor()
                    cur.execute("""
                        INSERT INTO incident_report (
                            stud_id, ir_date_incident, ir_description, ir_action_taken,
                            ir_referral, ir_follow_up_date, created_at, updated_at
                        ) VALUES (%s, %s, %s, %s, %s, %s, NOW(), NOW())
                    """, (lrn, date, desc, action, referral, followup))
                    conn.commit()
                    cur.close()
                    if not self.conn:
                        conn.close()
                    QMessageBox.information(self, "Success", "Report saved.")
                    self.load_data()
                except Exception as e:
                    if conn:
                        conn.rollback()
                    QMessageBox.critical(self, "Database Error", f"Could not save report: {e}")

    def edit_record(self):
        selected = self.table.currentRow()
        if selected < 0:
            QMessageBox.warning(self, "Edit Record", "Please select a row to edit.")
            return
        report_row = self.rows[selected]
        dialog = EditReportDialog(report_row, self)
        if dialog.exec_() == QDialog.Accepted:
            data = dialog.get_data()
            if data:
                lrn, date, desc, action, referral, followup = data

                def empty_to_none(x):
                    return x if x.strip() != "" else None

                lrn = lrn.strip()
                date = date.strip()
                desc = desc.strip()
                action = empty_to_none(action)
                referral = empty_to_none(referral)
                followup = empty_to_none(followup)

                try:
                    conn = self.get_conn()
                    cur = conn.cursor()
                    cur.execute("""
                        UPDATE incident_report
                        SET stud_id=%s, ir_date_incident=%s, ir_description=%s, ir_action_taken=%s,
                            ir_referral=%s, ir_follow_up_date=%s, updated_at=NOW()
                        WHERE ir_id=%s
                    """, (lrn, date, desc, action, referral, followup, report_row[0]))
                    conn.commit()
                    cur.close()
                    if not self.conn:
                        conn.close()
                    QMessageBox.information(self, "Success", "Report updated.")
                    self.load_data()
                except Exception as e:
                    if conn:
                        conn.rollback()
                    QMessageBox.critical(self, "Database Error", f"Could not update report: {e}")

    def delete_record(self):
        selected = self.table.currentRow()
        if selected < 0:
            QMessageBox.warning(self, "Delete Record", "Please select a row to delete.")
            return
        report_row = self.rows[selected]
        lrn = report_row[1]
        date = str(report_row[5])
        confirm = QMessageBox.question(
            self, "Delete Report",
            f"Delete report for Student LRN {lrn} on {date}?",
            QMessageBox.Yes | QMessageBox.No
        )
        if confirm == QMessageBox.Yes:
            try:
                conn = self.get_conn()
                cur = conn.cursor()
                cur.execute(
                    "DELETE FROM incident_report WHERE ir_id=%s",
                    (report_row[0],)
                )
                conn.commit()
                cur.close()
                if not self.conn:
                    conn.close()
                self.load_data()
            except Exception as e:
                if conn:
                    conn.rollback()
                QMessageBox.critical(self, "Database Error", f"Error deleting report:\n{e}")

    def view_record(self):
        selected = self.table.currentRow()
        if selected < 0:
            QMessageBox.warning(self, "View Record", "Please select a row to view.")
            return
        report_row = self.rows[selected]
        dialog = ViewReportDialog(report_row, self)
        dialog.exec_()

class AddReportDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Add Report")
        self.setMinimumWidth(550)
        self.setMinimumHeight(500)
        layout = QVBoxLayout()
        layout.setContentsMargins(30, 30, 30, 30)
        layout.setSpacing(20)
        self.setLayout(layout)
        self.setStyleSheet("""
            QDialog {
                background-color: #f6fbfc;
            }
            QLabel {
                font-size: 18px;
                font-weight: bold;
                padding: 5px;
                background-color: #f6fbfc;
            }
            QLineEdit, QDateEdit, QTextEdit {
                font-size: 18px;
                padding: 10px;
                border: 2px solid #D295BE;
                border-radius: 8px;
                min-height: 30px;
                background-color: #FFFFFF;
            }
            QPushButton {
                font-size: 16px;
                padding: 8px 16px;
            }
        """)

        form_layout = QFormLayout()
        form_layout.setLabelAlignment(Qt.AlignLeft | Qt.AlignVCenter)
        form_layout.setFormAlignment(Qt.AlignHCenter | Qt.AlignVCenter)
        form_layout.setSpacing(16)

        self.input_sid = QLineEdit()
        self.input_sid.setPlaceholderText("Enter Student LRN")
        form_layout.addRow("Student LRN", self.input_sid)

        self.input_date = QDateEdit(QDate.currentDate())
        self.input_date.setCalendarPopup(True)
        form_layout.addRow("Date of Incident", self.input_date)

        self.input_desc = QTextEdit()
        self.input_desc.setPlaceholderText("Describe the accident/incident")
        form_layout.addRow(QLabel("Description"), self.input_desc)

        self.input_action = QTextEdit()
        self.input_action.setPlaceholderText("Describe the action taken")
        form_layout.addRow(QLabel("Action Taken"), self.input_action)

        self.input_ref = QLineEdit()
        self.input_ref.setPlaceholderText("Enter referral (if any)")
        form_layout.addRow("Referral", self.input_ref)

        self.input_followup = QDateEdit(QDate.currentDate())
        self.input_followup.setCalendarPopup(True)
        form_layout.addRow("Follow-up Date", self.input_followup)

        layout.addLayout(form_layout)

        btn_row = QHBoxLayout()
        save_btn = QPushButton("Save Report")
        save_btn.setStyleSheet("""
            QPushButton {
                background:#D295BE; color:white; font-size: 17px; font-weight:bold;
                padding: 10px 26px; border-radius: 10px; min-width:160px;
            }
            QPushButton:hover { background: #2980b9; }
        """)
        save_btn.clicked.connect(self.save_and_accept)
        cancel_btn = QPushButton("Cancel")
        cancel_btn.setStyleSheet("""
            QPushButton {
                background:#b0b0b0; color:white; font-size: 17px;
                padding: 10px 26px; border-radius: 10px; min-width:120px;
            }
            QPushButton:hover { background: #888;}
        """)
        cancel_btn.clicked.connect(self.reject)
        btn_row.addWidget(save_btn)
        btn_row.addWidget(cancel_btn)
        layout.addLayout(btn_row)

        self.data = None

    def save_and_accept(self):
        lrn = self.input_sid.text().strip()
        date = self.input_date.date().toString("yyyy-MM-dd")
        desc = self.input_desc.toPlainText().strip()
        action = self.input_action.toPlainText().strip()
        referral = self.input_ref.text().strip()
        followup = self.input_followup.date().toString("yyyy-MM-dd")
        if not lrn or not lrn.isdigit() or not desc:
            QMessageBox.warning(self, "Missing Information", "Please provide a numeric LRN and description.")
            return
        self.data = (lrn, date, desc, action, referral, followup)
        self.accept()

    def get_data(self):
        return self.data

class EditReportDialog(QDialog):
    def __init__(self, report_row, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Edit Report")
        self.setMinimumWidth(550)
        self.setMinimumHeight(500)
        layout = QVBoxLayout()
        layout.setContentsMargins(30, 30, 30, 30)
        layout.setSpacing(20)
        self.setLayout(layout)
        self.setStyleSheet("""
            QDialog {
                background-color: #f6fbfc;
            }
            QLabel {
                font-size: 18px;
                font-weight: bold;
                padding: 5px;
                background-color: #f6fbfc;
            }
            QLineEdit, QDateEdit, QTextEdit {
                font-size: 18px;
                padding: 10px;
                border: 2px solid #D295BE;
                border-radius: 8px;
                min-height: 30px;
                background-color: #FFFFFF;
            }
            QPushButton {
                font-size: 16px;
                padding: 8px 16px;
            }
        """)

        form_layout = QFormLayout()
        form_layout.setLabelAlignment(Qt.AlignLeft | Qt.AlignVCenter)
        form_layout.setFormAlignment(Qt.AlignHCenter | Qt.AlignVCenter)
        form_layout.setSpacing(16)

        self.input_sid = QLineEdit(str(report_row[1] if report_row[1] else ""))
        self.input_sid.setPlaceholderText("Enter Student LRN")
        form_layout.addRow("Student LRN", self.input_sid)

        self.input_date = QDateEdit()
        self.input_date.setCalendarPopup(True)
        try:
            self.input_date.setDate(QDate.fromString(str(report_row[5]), "yyyy-MM-dd"))
        except:
            self.input_date.setDate(QDate.currentDate())
        form_layout.addRow("Date of Incident", self.input_date)

        self.input_desc = QTextEdit(str(report_row[6] if report_row[6] else ""))
        self.input_desc.setPlaceholderText("Describe the accident/incident")
        form_layout.addRow(QLabel("Description"), self.input_desc)

        self.input_action = QTextEdit(str(report_row[7] if report_row[7] else ""))
        self.input_action.setPlaceholderText("Describe the action taken")
        form_layout.addRow(QLabel("Action Taken"), self.input_action)

        self.input_ref = QLineEdit(str(report_row[8] if report_row[8] else ""))
        self.input_ref.setPlaceholderText("Enter referral (if any)")
        form_layout.addRow("Referral", self.input_ref)

        self.input_followup = QDateEdit()
        self.input_followup.setCalendarPopup(True)
        try:
            self.input_followup.setDate(QDate.fromString(str(report_row[9]), "yyyy-MM-dd"))
        except:
            self.input_followup.setDate(QDate.currentDate())
        form_layout.addRow("Follow-up Date", self.input_followup)

        layout.addLayout(form_layout)

        btn_row = QHBoxLayout()
        save_btn = QPushButton("Save Changes")
        save_btn.setStyleSheet("""
            QPushButton {
                background:#D295BE; color:white; font-size: 17px; font-weight:bold;
                padding: 10px 26px; border-radius: 10px; min-width:160px;
            }
            QPushButton:hover { background: #2980b9; }
        """)
        save_btn.clicked.connect(self.save_and_accept)
        cancel_btn = QPushButton("Cancel")
        cancel_btn.setStyleSheet("""
            QPushButton {
                background:#b0b0b0; color:white; font-size: 17px;
                padding: 10px 26px; border-radius: 10px; min-width:120px;
            }
            QPushButton:hover { background: #888;}
        """)
        cancel_btn.clicked.connect(self.reject)
        btn_row.addWidget(save_btn)
        btn_row.addWidget(cancel_btn)
        layout.addLayout(btn_row)

        self.data = None

    def save_and_accept(self):
        lrn = self.input_sid.text().strip()
        date = self.input_date.date().toString("yyyy-MM-dd")
        desc = self.input_desc.toPlainText().strip()
        action = self.input_action.toPlainText().strip()
        referral = self.input_ref.text().strip()
        followup = self.input_followup.date().toString("yyyy-MM-dd")
        if not lrn or not lrn.isdigit() or not desc:
            QMessageBox.warning(self, "Missing Information", "Please provide a numeric LRN and description.")
            return
        self.data = (lrn, date, desc, action, referral, followup)
        self.accept()

    def get_data(self):
        return self.data

class ViewReportDialog(QDialog):
    def __init__(self, report_row, parent=None):
        super().__init__(parent)
        self.setWindowTitle("View Report")
        self.setMinimumWidth(550)
        self.setMinimumHeight(500)
        layout = QVBoxLayout()
        layout.setContentsMargins(30, 30, 30, 30)
        layout.setSpacing(20)
        self.setLayout(layout)
        self.setStyleSheet("""
            QDialog {
                background-color: #f6fbfc;
            }
            QLabel {
                font-size: 18px;
                font-weight: bold;
                padding: 5px;
                background-color: #f6fbfc;
            }
            QLineEdit, QDateEdit, QTextEdit {
                font-size: 18px;
                padding: 10px;
                border: 2px solid #D295BE;
                border-radius: 8px;
                min-height: 30px;
                background-color: #F9F9F9;
            }
            QLineEdit:disabled, QDateEdit:disabled, QTextEdit:disabled {
                color: #888;
            }
            QPushButton {
                font-size: 16px;
                padding: 8px 16px;
            }
        """)

        form_layout = QFormLayout()
        form_layout.setLabelAlignment(Qt.AlignLeft | Qt.AlignVCenter)
        form_layout.setFormAlignment(Qt.AlignHCenter | Qt.AlignVCenter)
        form_layout.setSpacing(16)

        self.input_sid = QLineEdit(str(report_row[1] if report_row[1] else ""))
        self.input_sid.setReadOnly(True)
        form_layout.addRow("Student LRN", self.input_sid)

        self.input_fname = QLineEdit(str(report_row[2] if report_row[2] else ""))
        self.input_fname.setReadOnly(True)
        form_layout.addRow("First Name", self.input_fname)

        self.input_lname = QLineEdit(str(report_row[3] if report_row[3] else ""))
        self.input_lname.setReadOnly(True)
        form_layout.addRow("Last Name", self.input_lname)

        self.input_mname = QLineEdit(str(report_row[4] if report_row[4] else ""))
        self.input_mname.setReadOnly(True)
        form_layout.addRow("Middle Name", self.input_mname)

        self.input_date = QDateEdit()
        self.input_date.setCalendarPopup(True)
        try:
            self.input_date.setDate(QDate.fromString(str(report_row[5]), "yyyy-MM-dd"))
        except:
            self.input_date.setDate(QDate.currentDate())
        self.input_date.setReadOnly(True)
        self.input_date.setButtonSymbols(QDateEdit.NoButtons)
        form_layout.addRow("Date of Incident", self.input_date)

        self.input_desc = QTextEdit(str(report_row[6] if report_row[6] else ""))
        self.input_desc.setReadOnly(True)
        form_layout.addRow(QLabel("Description"), self.input_desc)

        self.input_action = QTextEdit(str(report_row[7] if report_row[7] else ""))
        self.input_action.setReadOnly(True)
        form_layout.addRow(QLabel("Action Taken"), self.input_action)

        self.input_ref = QLineEdit(str(report_row[8] if report_row[8] else ""))
        self.input_ref.setReadOnly(True)
        form_layout.addRow("Referral", self.input_ref)

        self.input_followup = QDateEdit()
        self.input_followup.setCalendarPopup(True)
        try:
            self.input_followup.setDate(QDate.fromString(str(report_row[9]), "yyyy-MM-dd"))
        except:
            self.input_followup.setDate(QDate.currentDate())
        self.input_followup.setReadOnly(True)
        self.input_followup.setButtonSymbols(QDateEdit.NoButtons)
        form_layout.addRow("Follow-up Date", self.input_followup)

        self.input_created_at = QLineEdit(str(report_row[10] if report_row[10] else ""))
        self.input_created_at.setReadOnly(True)
        form_layout.addRow("Created At", self.input_created_at)

        layout.addLayout(form_layout)

        btn_row = QHBoxLayout()
        close_btn = QPushButton("Close")
        close_btn.setStyleSheet("""
            QPushButton {
                background:#b0b0b0; color:white; font-size: 17px;
                padding: 10px 26px; border-radius: 10px; min-width:120px;
            }
            QPushButton:hover { background: #888;}
        """)
        close_btn.clicked.connect(self.reject)
        btn_row.addStretch(1)
        btn_row.addWidget(close_btn)
        layout.addLayout(btn_row)
