from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QLabel, QPushButton, QTableWidget, QTableWidgetItem,
    QHBoxLayout, QLineEdit, QTextEdit, QDateEdit, QMessageBox, QFormLayout, QFileDialog,
    QStackedWidget, QHeaderView, QDialog, QComboBox
)
from PyQt5.QtCore import Qt, QDate, pyqtSignal
import psycopg2
from docx import Document
from datetime import datetime

def get_db_connection():
    return psycopg2.connect(
        host="localhost",
        database="MediTrackSNHS",
        user="postgres",
        password="Mylovemondejar"
    )

class ReferralPage(QWidget):
    def __init__(self):
        super().__init__()
        self.setObjectName("ReferralPageBg")
        self.records_page = ReferralRecordsPage()
        layout = QVBoxLayout(self)
        layout.addWidget(self.records_page)
        self.setLayout(layout)
        self.setStyleSheet("""
            QWidget#ReferralPageBg {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #FED6C8, stop:1 #F9E6DC);
            }
        """)

class ReferralRecordsPage(QWidget):
    def __init__(self):
        super().__init__()
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 60)
        layout.setSpacing(10)

        title = QLabel("Student Referrals")
        title.setStyleSheet("""
            font-size: 28px;
            font-weight: bold;
            background: #FDD1B0;
            border-radius: 10px;
            padding: 10px;
            color: #3a2b23;
        """)
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)

        control_layout = QHBoxLayout()
        self.search_bar = QLineEdit()
        self.search_bar.setFixedHeight(45)
        self.search_bar.setFixedWidth(400)
        self.search_bar.setStyleSheet("""
            QLineEdit {
                background-color: #FFF0F5;
                border: 2px solid #D295BF;
                border-radius: 5px;
                padding: 8px 12px;
                font-size: 16px;
                color: #333;
            }
            QLineEdit:focus {
                border: 2px solid #BA74AD;
                background-color: #FFE9F1;
            }
        """)
        self.search_bar.setPlaceholderText("Search referral...")
        self.search_bar.textChanged.connect(self.search_referrals)

        add_btn = QPushButton("Add Referral")
        add_btn.setFixedWidth(150)
        add_btn.setStyleSheet("background-color: #D295BF; color: white; font-weight: bold; padding: 15px; border-radius: 10px; font-size: 18px;")
        add_btn.clicked.connect(self.open_add_referral_dialog)
        
        self.edit_btn = QPushButton("Edit")
        self.edit_btn.setFixedWidth(120)
        self.edit_btn.setStyleSheet("background-color: #F7CA18; color: white; font-weight: bold; padding: 15px; border-radius: 10px; font-size: 18px;")
        self.edit_btn.clicked.connect(self.edit_record)
        
        self.view_btn = QPushButton("View")
        self.view_btn.setFixedWidth(120)
        self.view_btn.setStyleSheet("background-color: #5DADE2; color: white; font-weight: bold; padding: 15px; border-radius: 10px; font-size: 18px;")
        self.view_btn.clicked.connect(self.view_record)
        
        self.delete_btn = QPushButton("Delete")
        self.delete_btn.setFixedWidth(120)
        self.delete_btn.setStyleSheet("background-color: #FF9999; color: white; font-weight: bold; padding: 15px; border-radius: 10px; font-size: 18px;")
        self.delete_btn.clicked.connect(self.delete_record)

        control_layout.addWidget(self.search_bar)
        control_layout.addStretch()
        control_layout.addWidget(add_btn)
        control_layout.addWidget(self.edit_btn)
        control_layout.addWidget(self.view_btn)
        control_layout.addWidget(self.delete_btn)
        layout.addLayout(control_layout)

        self.table = QTableWidget()
        self.table.setColumnCount(8)
        self.table.setHorizontalHeaderLabels(["Student ID", "Student Name", "Date", "Reason", "Referred To", "Status", "Notes", "Actions"])
        self.table.setStyleSheet("""
            QTableWidget {
                background-color: #FFF5F0;
            }
            QTableWidget::item {
                font-size: 18px;
            }
            QHeaderView::section {
                font-size: 20px;
                font-weight: bold;
                background-color: #FFE1D3;
                border: 1px solid #D295BF;
            }
            QTableCornerButton::section {
                background-color: #FFE1D3;
                border: 1px solid #D295BF;
            }
        """)
        self.table.setSelectionBehavior(QTableWidget.SelectRows)
        self.table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.table.setSortingEnabled(False)
        header = self.table.horizontalHeader()
        for i in range(self.table.columnCount()):
            header.setSectionResizeMode(i, QHeaderView.Stretch)
        self.table.horizontalHeader().setStretchLastSection(True)
        self.setLayout(layout)
        layout.addWidget(self.table)
        self.load_data()

    def open_add_referral_dialog(self):
        dialog = AddReferralDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            data = dialog.get_data()
            if data:
                sid, name, date, reason, to, status, notes = data
                try:
                    conn = get_db_connection()
                    cur = conn.cursor()
                    cur.execute("""
                        INSERT INTO referral (stud_id, referral_date, referral_reason, referral_to, referral_status, notes, created_at)
                        VALUES (%s, %s, %s, %s, %s, %s, NOW())
                    """, (sid, date, reason, to, status, notes))
                    conn.commit()
                    cur.close()
                    conn.close()
                    QMessageBox.information(self, "Success", "Referral saved.")
                    self.load_data()
                except Exception as e:
                    QMessageBox.critical(self, "Database Error", f"Could not save referral: {e}")

    def load_data(self):
        self.table.setRowCount(0)
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("""
                SELECT r.referral_id, r.stud_id, 
                       s.stud_lname || ', ' || s.stud_fname ||
                       CASE WHEN s.stud_mname IS NOT NULL AND s.stud_mname <> '' THEN ' ' || s.stud_mname ELSE '' END AS stud_fullname,
                       r.referral_date, r.referral_reason, r.referral_to, r.referral_status, r.notes
                FROM referral r
                LEFT JOIN student s ON r.stud_id = s.stud_id
                ORDER BY r.referral_date DESC
            """)
            rows = cur.fetchall()
            self.rows = rows  # for searching/filtering
            self.display_rows(self.rows)
            cur.close()
            conn.close()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load referrals: {e}")

    def display_rows(self, rows):
        self.table.setRowCount(len(rows))
        for row_idx, row in enumerate(rows):
            for col_idx, val in enumerate(row[1:]):
                item = QTableWidgetItem(str(val) if val is not None else "")
                font = item.font()
                font.setPointSize(9)
                item.setFont(font)
                self.table.setItem(row_idx, col_idx, item)
            self.table.setVerticalHeaderItem(row_idx, QTableWidgetItem(str(row_idx + 1)))
            btn = QPushButton("Generate DOCX")
            btn.setFixedWidth(150)
            btn.setStyleSheet("""
                QPushButton { background: #A3D5FF; border-radius:5px; font-size:16px; padding:6px 12px;}
            """)
            btn.clicked.connect(lambda _, r=row: self.generate_referral_docx(r))
            btn_layout = QHBoxLayout()
            btn_layout.addWidget(btn)
            btn_layout.setContentsMargins(0,0,0,0)
            btn_widget = QWidget()
            btn_widget.setStyleSheet("background-color: #FFF5F0;")
            btn_widget.setLayout(btn_layout)
            self.table.setCellWidget(row_idx, 7, btn_widget)

    def search_referrals(self):
        query = self.search_bar.text().lower()
        filtered = []
        for row in self.rows:
            # row = (referral_id, stud_id, stud_fullname, referral_date, referral_reason, referral_to, referral_status, notes)
            if (query in str(row[1]).lower() or    # Student ID
                query in str(row[2]).lower() or    # Student Name
                query in str(row[3]).lower() or    # Referral Date
                query in str(row[4]).lower() or    # Reason
                query in str(row[5]).lower() or    # Referral To
                query in str(row[6]).lower() or    # Status
                query in str(row[7]).lower()):     # Notes
                filtered.append(row)
        self.display_rows(filtered)

    def generate_referral_docx(self, referral_row):
        referral = {
            "referral_id": referral_row[0],
            "stud_id": referral_row[1],
            "stud_fullname": referral_row[2],
            "referral_date": str(referral_row[3]),
            "referral_reason": referral_row[4],
            "referral_to": referral_row[5],
            "referral_status": referral_row[6],
            "notes": referral_row[7]
        }
        # Fetch student info for year level and section
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("""
                SELECT year_level_id, section_id
                FROM student
                WHERE stud_id = %s
            """, (referral["stud_id"],))
            stu = cur.fetchone()
            year_level_name = ""
            section_name = ""
            if stu:
                year_level_id, section_id = stu
                if year_level_id:
                    cur.execute("SELECT year_level_name FROM year_level WHERE year_level_id=%s", (year_level_id,))
                    yl = cur.fetchone()
                    year_level_name = yl[0] if yl else ""
                if section_id:
                    cur.execute("SELECT section_name FROM section WHERE section_id=%s", (section_id,))
                    sc = cur.fetchone()
                    section_name = sc[0] if sc else ""
            cur.close()
            conn.close()
        except Exception as e:
            QMessageBox.critical(self, "Database Error", f"Could not fetch student info: {e}")
            return

        default_filename = f"Referral_{referral['stud_id']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        save_path, _ = QFileDialog.getSaveFileName(self, "Save Referral DOCX", default_filename, "Word Documents (*.docx)")
        if not save_path:
            return

        try:
            doc = Document()
            doc.add_heading('School Clinic Referral Form', 0)
            doc.add_paragraph(f"Date: {referral['referral_date']}")
            doc.add_paragraph(f"Student Name: {referral['stud_fullname']}")
            doc.add_paragraph(f"Student ID: {referral['stud_id']}")
            doc.add_paragraph(f"Year Level: {year_level_name}")
            doc.add_paragraph(f"Section: {section_name}")
            doc.add_paragraph(f"Referral To: {referral['referral_to']}")
            doc.add_paragraph(f"Reason for Referral: {referral['referral_reason']}")
            doc.add_paragraph(f"Status: {referral['referral_status']}")
            doc.add_paragraph(f"Notes: {referral['notes']}")
            doc.add_paragraph("\nSignature: ___________________________")
            doc.add_paragraph("Clinic Staff")

            doc.save(save_path)
            QMessageBox.information(self, "Referral Doc Created", f"Referral document saved as:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "DOCX Error", f"Could not generate referral document: {e}")
            
    def edit_record(self):
        selected = self.table.currentRow()
        if selected < 0:
            QMessageBox.warning(self, "Edit Record", "Please select a row to edit.")
            return
        referral_row = self.rows[selected]
        dialog = EditReferralDialog(referral_row, self)
        if dialog.exec_() == QDialog.Accepted:
            data = dialog.get_data()
            if data:
                sid, name, date, reason, to, status, notes = data
                try:
                    conn = get_db_connection()
                    cur = conn.cursor()
                    # Update using referral_id (referral_row[0])
                    cur.execute("""
                        UPDATE referral
                        SET stud_id=%s, referral_date=%s, referral_reason=%s, referral_to=%s, referral_status=%s, notes=%s
                        WHERE referral_id=%s
                    """, (sid, date, reason, to, status, notes, referral_row[0]))
                    conn.commit()
                    cur.close()
                    conn.close()
                    QMessageBox.information(self, "Success", "Referral updated.")
                    self.load_data()
                except Exception as e:
                    QMessageBox.critical(self, "Database Error", f"Could not update referral: {e}")

    def delete_record(self):
        selected = self.table.currentRow()
        if selected < 0:
            QMessageBox.warning(self, "Delete Record", "Please select a row to delete.")
            return
        referral_row = self.rows[selected]
        stud_id = referral_row[1]
        referral_date = str(referral_row[3])
        confirm = QMessageBox.question(
            self, "Delete Referral",
            f"Delete referral for Student ID {stud_id} on {referral_date}?",
            QMessageBox.Yes | QMessageBox.No
        )
        if confirm == QMessageBox.Yes:
            try:
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute(
                    "DELETE FROM referral WHERE referral_id=%s",
                    (referral_row[0],)
                )
                conn.commit()
                cur.close()
                conn.close()
                self.load_data()
            except Exception as e:
                QMessageBox.critical(self, "Database Error", f"Error deleting referral:\n{e}")
    def view_record(self):
        selected = self.table.currentRow()
        if selected < 0:
            QMessageBox.warning(self, "View Record", "Please select a row to view.")
            return
        referral_row = self.rows[selected]
        dialog = ViewReferralDialog(referral_row, self)
        dialog.exec_()

class AddReferralDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Add Referral")
        self.setMinimumWidth(550)
        self.setMinimumHeight(500)  # or use self.resize(600, 500)
        layout = QVBoxLayout()
        layout.setContentsMargins(30, 30, 30, 30)  # left, top, right, bottom
        layout.setSpacing(20)  # more space between widgets
        self.setLayout(layout)
        self.setStyleSheet("""
            QDialog {
                background-color: #FFF5F0;
            }
            QLabel {
                font-size: 18px;
                font-weight: bold;
                padding: 5px;
                background-color: #FFF5F0;
            }
            QLineEdit, QComboBox, QDateEdit, QTextEdit {
                font-size: 18px;
                padding: 10px;
                border: 2px solid #D295BF;
                border-radius: 8px;
                min-height: 30px;
                background-color: #FFFFFF;
            }
            QPushButton {
                font-size: 16px;
                padding: 8px 16px;
            }
        """)


        form_layout = QFormLayout()
        form_layout.setLabelAlignment(Qt.AlignLeft | Qt.AlignVCenter)
        form_layout.setFormAlignment(Qt.AlignHCenter | Qt.AlignVCenter)
        form_layout.setSpacing(16)

        self.stud_id = QLineEdit()
        self.stud_id.setPlaceholderText("Enter Student LRN/ID")
        form_layout.addRow("Student ID", self.stud_id)

        self.stud_name = QLineEdit()
        self.stud_name.setPlaceholderText("Enter Student Name (optional)")
        form_layout.addRow(QLabel("Student Name"), self.stud_name)

        self.referral_date = QDateEdit(QDate.currentDate())
        self.referral_date.setCalendarPopup(True)
        form_layout.addRow("Date", self.referral_date)

        self.referral_reason = QLineEdit()
        self.referral_reason.setPlaceholderText("Enter referral reason")
        form_layout.addRow("Reason", self.referral_reason)

        self.referral_to = QLineEdit()
        self.referral_to.setPlaceholderText("Referred to whom?")
        form_layout.addRow("Referral To", self.referral_to)

        self.referral_status = QLineEdit()
        self.referral_status.setPlaceholderText("Enter referral status")
        form_layout.addRow("Status", self.referral_status)

        self.notes = QTextEdit()
        self.notes.setPlaceholderText("Any notes (optional)")
        form_layout.addRow(QLabel("Notes"), self.notes)

        layout.addLayout(form_layout)

        btn_row = QHBoxLayout()
        save_btn = QPushButton("Save Referral")
        save_btn.setStyleSheet("""
            QPushButton {
                background:#BA74AD; color:white; font-size: 17px; font-weight:bold;
                padding: 10px 26px; border-radius: 10px; min-width:160px;
            }
            QPushButton:hover { background: #d295bf; }
        """)
        save_btn.clicked.connect(self.save_and_accept)
        cancel_btn = QPushButton("Cancel")
        cancel_btn.setStyleSheet("""
            QPushButton {
                background:#b0b0b0; color:white; font-size: 17px;
                padding: 10px 26px; border-radius: 10px; min-width:120px;
            }
            QPushButton:hover { background: #888;}
        """)
        cancel_btn.clicked.connect(self.reject)
        btn_row.addWidget(save_btn)
        btn_row.addWidget(cancel_btn)
        layout.addLayout(btn_row)

        self.data = None
    def save_and_accept(self):
        sid = self.stud_id.text().strip()
        name = self.stud_name.text().strip()
        date = self.referral_date.date().toString("yyyy-MM-dd")
        reason = self.referral_reason.text().strip()
        to = self.referral_to.text().strip()
        status = self.referral_status.text().strip()
        notes = self.notes.toPlainText().strip()
        if not sid or not date or not reason or not to or not status:
            QMessageBox.warning(self, "Missing Information", "Please fill all required fields.")
            return
        self.data = (sid, name, date, reason, to, status, notes)
        self.accept()

    def get_data(self):
        return self.data

class EditReferralDialog(QDialog):
    def __init__(self, referral_row, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Edit Referral")
        self.setMinimumWidth(550)
        self.setMinimumHeight(500)
        layout = QVBoxLayout()
        layout.setContentsMargins(30, 30, 30, 30)
        layout.setSpacing(20)
        self.setLayout(layout)
        self.setStyleSheet("""
            QDialog {
                background-color: #FFF5F0;
            }
            QLabel {
                font-size: 18px;
                font-weight: bold;
                padding: 5px;
                background-color: #FFF5F0;
            }
            QLineEdit, QComboBox, QDateEdit, QTextEdit {
                font-size: 18px;
                padding: 10px;
                border: 2px solid #D295BF;
                border-radius: 8px;
                min-height: 30px;
                background-color: #FFFFFF;
            }
            QPushButton {
                font-size: 16px;
                padding: 8px 16px;
            }
        """)

        form_layout = QFormLayout()
        form_layout.setLabelAlignment(Qt.AlignLeft | Qt.AlignVCenter)
        form_layout.setFormAlignment(Qt.AlignHCenter | Qt.AlignVCenter)
        form_layout.setSpacing(16)

        self.stud_id = QLineEdit(str(referral_row[1] if referral_row[1] else ""))
        self.stud_id.setPlaceholderText("Enter Student LRN/ID")
        form_layout.addRow("Student ID", self.stud_id)

        self.stud_name = QLineEdit(str(referral_row[2] if referral_row[2] else ""))
        self.stud_name.setPlaceholderText("Enter Student Name (optional)")
        form_layout.addRow(QLabel("Student Name"), self.stud_name)

        self.referral_date = QDateEdit()
        self.referral_date.setCalendarPopup(True)
        # Set date, fallback to today if invalid
        try:
            self.referral_date.setDate(QDate.fromString(str(referral_row[3]), "yyyy-MM-dd"))
        except:
            self.referral_date.setDate(QDate.currentDate())
        form_layout.addRow("Date", self.referral_date)

        self.referral_reason = QLineEdit(str(referral_row[4] if referral_row[4] else ""))
        self.referral_reason.setPlaceholderText("Enter referral reason")
        form_layout.addRow("Reason", self.referral_reason)

        self.referral_to = QLineEdit(str(referral_row[5] if referral_row[5] else ""))
        self.referral_to.setPlaceholderText("Referred to whom?")
        form_layout.addRow("Referral To", self.referral_to)

        self.referral_status = QLineEdit(str(referral_row[6] if referral_row[6] else ""))
        self.referral_status.setPlaceholderText("Enter referral status")
        form_layout.addRow("Status", self.referral_status)

        self.notes = QTextEdit(str(referral_row[7] if referral_row[7] else ""))
        self.notes.setPlaceholderText("Any notes (optional)")
        form_layout.addRow(QLabel("Notes"), self.notes)

        layout.addLayout(form_layout)

        btn_row = QHBoxLayout()
        save_btn = QPushButton("Save Changes")
        save_btn.setStyleSheet("""
            QPushButton {
                background:#BA74AD; color:white; font-size: 17px; font-weight:bold;
                padding: 10px 26px; border-radius: 10px; min-width:160px;
            }
            QPushButton:hover { background: #d295bf; }
        """)
        save_btn.clicked.connect(self.save_and_accept)
        cancel_btn = QPushButton("Cancel")
        cancel_btn.setStyleSheet("""
            QPushButton {
                background:#b0b0b0; color:white; font-size: 17px;
                padding: 10px 26px; border-radius: 10px; min-width:120px;
            }
            QPushButton:hover { background: #888;}
        """)
        cancel_btn.clicked.connect(self.reject)
        btn_row.addWidget(save_btn)
        btn_row.addWidget(cancel_btn)
        layout.addLayout(btn_row)

        self.data = None

    def save_and_accept(self):
        sid = self.stud_id.text().strip()
        name = self.stud_name.text().strip()
        date = self.referral_date.date().toString("yyyy-MM-dd")
        reason = self.referral_reason.text().strip()
        to = self.referral_to.text().strip()
        status = self.referral_status.text().strip()
        notes = self.notes.toPlainText().strip()
        if not sid or not date or not reason or not to or not status:
            QMessageBox.warning(self, "Missing Information", "Please fill all required fields.")
            return
        self.data = (sid, name, date, reason, to, status, notes)
        self.accept()

    def get_data(self):
        return self.data

class ViewReferralDialog(QDialog):
    def __init__(self, referral_row, parent=None):
        super().__init__(parent)
        self.setWindowTitle("View Referral")
        self.setMinimumWidth(550)
        self.setMinimumHeight(500)
        layout = QVBoxLayout()
        layout.setContentsMargins(30, 30, 30, 30)
        layout.setSpacing(20)
        self.setLayout(layout)
        self.setStyleSheet("""
            QDialog {
                background-color: #FFF5F0;
            }
            QLabel {
                font-size: 18px;
                font-weight: bold;
                padding: 5px;
                background-color: #FFF5F0;
            }
            QLineEdit, QComboBox, QDateEdit, QTextEdit {
                font-size: 18px;
                padding: 10px;
                border: 2px solid #D295BF;
                border-radius: 8px;
                min-height: 30px;
                background-color: #F9F9F9;
            }
            QLineEdit:disabled, QComboBox:disabled, QDateEdit:disabled, QTextEdit:disabled {
                color: #888;
            }
            QPushButton {
                font-size: 16px;
                padding: 8px 16px;
            }
        """)

        form_layout = QFormLayout()
        form_layout.setLabelAlignment(Qt.AlignLeft | Qt.AlignVCenter)
        form_layout.setFormAlignment(Qt.AlignHCenter | Qt.AlignVCenter)
        form_layout.setSpacing(16)

        self.stud_id = QLineEdit(str(referral_row[1] if referral_row[1] else ""))
        self.stud_id.setReadOnly(True)
        form_layout.addRow("Student ID", self.stud_id)

        self.stud_name = QLineEdit(str(referral_row[2] if referral_row[2] else ""))
        self.stud_name.setReadOnly(True)
        form_layout.addRow(QLabel("Student Name"), self.stud_name)

        self.referral_date = QDateEdit()
        self.referral_date.setCalendarPopup(True)
        try:
            self.referral_date.setDate(QDate.fromString(str(referral_row[3]), "yyyy-MM-dd"))
        except:
            self.referral_date.setDate(QDate.currentDate())
        self.referral_date.setReadOnly(True)
        self.referral_date.setButtonSymbols(QDateEdit.NoButtons)
        form_layout.addRow("Date", self.referral_date)

        self.referral_reason = QLineEdit(str(referral_row[4] if referral_row[4] else ""))
        self.referral_reason.setReadOnly(True)
        form_layout.addRow("Reason", self.referral_reason)

        self.referral_to = QLineEdit(str(referral_row[5] if referral_row[5] else ""))
        self.referral_to.setReadOnly(True)
        form_layout.addRow("Referral To", self.referral_to)

        self.referral_status = QLineEdit(str(referral_row[6] if referral_row[6] else ""))
        self.referral_status.setReadOnly(True)
        form_layout.addRow("Status", self.referral_status)

        self.notes = QTextEdit(str(referral_row[7] if referral_row[7] else ""))
        self.notes.setReadOnly(True)
        form_layout.addRow(QLabel("Notes"), self.notes)

        layout.addLayout(form_layout)

        btn_row = QHBoxLayout()
        close_btn = QPushButton("Close")
        close_btn.setStyleSheet("""
            QPushButton {
                background:#b0b0b0; color:white; font-size: 17px;
                padding: 10px 26px; border-radius: 10px; min-width:120px;
            }
            QPushButton:hover { background: #888;}
        """)
        close_btn.clicked.connect(self.reject)
        btn_row.addStretch(1)
        btn_row.addWidget(close_btn)
        layout.addLayout(btn_row)
